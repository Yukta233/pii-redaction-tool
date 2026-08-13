"""
Redacts PII inside a .docx file while preserving formatting/layout.

Strategy: python-docx exposes text at the *paragraph* level (concatenating
all runs). We redact at the paragraph level (and separately for each table
cell's paragraphs) to catch PII that Word has split across multiple runs,
then write the result back into the paragraph's first run and clear the
rest. This trades exact run-level formatting granularity (e.g. a name that
was half-bold/half-not) for reliably catching PII that spans run
boundaries, which is the more common failure mode in real Word documents.
"""

from docx import Document
from .redactor import Redactor


def _redact_paragraph(paragraph, redactor: Redactor, full_log: list):
    original = paragraph.text
    if not original.strip():
        return
    new_text, log = redactor.redact(original)
    if new_text == original:
        return
    full_log.extend(log)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def redact_docx(input_path: str, output_path: str) -> list:
    doc = Document(input_path)
    redactor = Redactor()
    full_log = []

    for paragraph in doc.paragraphs:
        _redact_paragraph(paragraph, redactor, full_log)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _redact_paragraph(paragraph, redactor, full_log)

    # headers / footers
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                _redact_paragraph(paragraph, redactor, full_log)

    doc.save(output_path)
    return full_log
